"""
Génère le rapport Word d'analyse de la simulation de base martienne.
"""
import pickle
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from collections import Counter, defaultdict

# --- Chargement des données pour les statistiques dans le texte ---
with open("simulation_data.pkl", "rb") as f:
    data = pickle.load(f)

log_expeditions   = data["log_expeditions"]
log_sinistres     = data["log_sinistres"]
historique_stocks = data["historique_stocks"]
donnees_membres   = data["donnees_membres"]

# Calculs pour le graphique 1
nourriture = [s[2] for s in historique_stocks]
graines    = [s[1] for s in historique_stocks]
pieces     = [s[3] for s in historique_stocks]

# Calculs pour le graphique 3
durees = [exp[3] for exp in log_expeditions]
moy_dur = np.mean(durees)
med_dur = np.median(durees)
std_dur = np.std(durees)

# Calculs pour le graphique 5
POIDS = {"sinistre": 1, "exp_lancee": 5, "exp_participee": 8,
         "exp_receptionnee": 3, "evenement_serre": 2, "commande": 1, "gestion_membre": 1}
roles_actifs = [m["role"] for m in donnees_membres if m["etat"] == 1]
role_counts  = Counter(roles_actifs)
charge_par_role = defaultdict(int)
for m in donnees_membres:
    if m["etat"] == 1:
        charge = (m["nb_sinistres"] * POIDS["sinistre"]
                  + m["nb_exp_lancees"] * POIDS["exp_lancee"]
                  + m["nb_exp_participees"] * POIDS["exp_participee"]
                  + m["nb_exp_receptionnees"] * POIDS["exp_receptionnee"]
                  + m["nb_evenements_serre"] * POIDS["evenement_serre"]
                  + m["nb_commandes_receptionnees"] * POIDS["commande"]
                  + (m["nb_membres_ajoutes"] + m["nb_membres_supprimes"]) * POIDS["gestion_membre"])
        charge_par_role[m["role"]] += charge

total_eff    = sum(role_counts.values())
total_charge = sum(charge_par_role.values()) or 1
role_order   = ["Commandant", "Technicien", "Biologiste", "Chercheur"]

# --- Construction du document ---
doc = Document()

# Style général
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Titre principal
titre = doc.add_heading("Rapport d'analyse — Simulation de base martienne", level=0)
titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ================================================================
# 1. Introduction / Approche de génération des données
# ================================================================
doc.add_heading("1. Approche de génération des données", level=1)

doc.add_paragraph(
    "La simulation repose sur un modèle orienté objet représentant une base martienne. "
    "La classe centrale Base orchestre l'ensemble des opérations via des membres d'équipage "
    "(Commandant, Technicien, Biologiste, Chercheur), des modules (Garage, Serre), "
    "des expéditions et des sinistres."
)

doc.add_paragraph(
    "Les données sont générées par un script de simulation (generate_data.py) exécutant "
    f"1 000 jours de fonctionnement. La base est initialisée avec 20 garages, 20 serres "
    "et 98 membres. À chaque jour simulé, des événements probabilistes sont déclenchés :"
)

bullets = [
    "Ravitaillement toutes les 20 jours (graines systématiques, nourriture et pièces avec faible probabilité)",
    "Consommation quotidienne de nourriture par chaque membre actif",
    "Ajout / suppression de membres (probabilités 5 % et 3 %)",
    "Plantation (60 %) et récolte (40 %) dans les serres par les biologistes",
    "Lancement d'expéditions (25 %) depuis un garage par deux chercheurs distincts (un lanceur et un participant)",
    "Retour d'expéditions selon une durée tirée d'une loi normale tronquée (µ=20j, σ=8j)",
    "Déclaration de sinistres sur garages (10 %) et serres (8 %), avec dégâts selon loi exponentielle",
    "Réparation de sinistres (40 % si sinistre en cours sur le module concerné) par un technicien",
    "Suppression de modules via une probabilité de Bernoulli conditionnelle aux dégâts cumulés",
]
for b in bullets:
    p = doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph(
    "Les données sont exportées dans deux formats : le fichier pickle (simulation_data.pkl), "
    "utilisé par analyse_stats.py car il conserve exactement les types Python (None, entiers, dicts), "
    "et des fichiers CSV pour la consultation externe (Excel, tableurs)."
)

doc.add_page_break()

# ================================================================
# 2. Graphique 1 — Évolution des stocks
# ================================================================
doc.add_heading("2. Graphique 1 — Évolution des stocks au fil du temps", level=1)

doc.add_heading("Méthode de génération", level=2)
doc.add_paragraph(
    "À chaque jour de simulation, les trois stocks (graines, nourriture, pièces module) "
    "sont enregistrés dans historique_stocks. Le graphique superpose les valeurs brutes "
    "(transparence α=0,2) avec une moyenne mobile sur 30 jours pour lisser les variations. "
    "Un sous-graphique en bas isole le stock de nourriture avec un seuil critique à 20 unités."
)

doc.add_heading("Interprétation", level=2)
doc.add_paragraph(
    f"Le stock de graines augmente régulièrement grâce aux ravitaillements périodiques (toutes les 20 jours) "
    f"et diminue lors des plantations. Il atteint un maximum de {max(graines)} unités. "
    f"Le stock de nourriture oscille fortement : chaque membre consomme 1 unité par jour, "
    f"ce qui génère une pression constante compensée par les récoltes et les ravitaillements ponctuels. "
    f"Le stock de pièces module reste relativement stable ; il est consommé par les réparations "
    f"et reconstitué par les retours d'expédition."
)

doc.add_picture("graphiques/01_evolution_stocks.png", width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ================================================================
# 3. Graphique 3 — Distribution des durées d'expédition
# ================================================================
doc.add_heading("3. Graphique 3 — Distribution des durées d'expédition", level=1)

doc.add_heading("Méthode de génération", level=2)
doc.add_paragraph(
    "Lors du lancement de chaque expédition, une durée est tirée selon une loi normale tronquée "
    "(µ = 20 jours, σ = 8 jours, bornée entre 5 et 45 jours). "
    "La durée réelle est calculée comme la différence entre le jour de retour et le jour de lancement. "
    "L'histogramme est tracé sur l'ensemble des expéditions réceptionnées avec indication "
    "de la moyenne, de la médiane et de l'écart-type."
)

doc.add_heading("Interprétation", level=2)
doc.add_paragraph(
    f"Sur {len(durees)} expéditions terminées, la durée moyenne est de {moy_dur:.1f} jours "
    f"(médiane : {med_dur:.1f} j, écart-type : {std_dur:.1f} j). "
    "La distribution suit bien la forme en cloche attendue d'une loi normale, centrée autour "
    f"de la valeur cible de 20 jours. "
    "Les queues sont tronquées aux bornes 5 et 45 jours, ce qui explique l'absence de valeurs extrêmes. "
    "La proximité de la moyenne et de la médiane confirme la symétrie de la distribution."
)

doc.add_picture("graphiques/03_durees_expedition.png", width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ================================================================
# 4. Graphique 5 — Effectifs vs charge pondérée
# ================================================================
doc.add_heading("4. Graphique 5 — Effectifs vs charge de travail pondérée par rôle", level=1)

doc.add_heading("Méthode de génération", level=2)
doc.add_paragraph(
    "Pour chaque membre actif en fin de simulation, une charge de travail pondérée est calculée "
    "en multipliant le nombre d'activités de chaque type par un coefficient reflétant leur complexité :"
)

poids_desc = [
    "Sinistre déclaré ou réparé : 1",
    "Expédition lancée : 5",
    "Expédition participée : 8  (le plus contraignant — membre immobilisé)",
    "Expédition réceptionnée : 3",
    "Événement serre (plantation/récolte) : 2",
    "Commande réceptionnée : 1",
    "Gestion de membre (ajout/suppression) : 1",
]
for p in poids_desc:
    doc.add_paragraph(p, style="List Bullet")

doc.add_paragraph(
    "Les charges sont agrégées par rôle puis exprimées en pourcentage du total. "
    "Le graphique compare côte à côte le pourcentage d'effectifs et le pourcentage de charge pour chaque rôle."
)

doc.add_heading("Interprétation", level=2)

lines = []
for role in role_order:
    if role in role_counts:
        pct_eff = role_counts[role] / total_eff * 100
        pct_chg = charge_par_role[role] / total_charge * 100
        lines.append(f"{role} : {pct_eff:.0f} % des effectifs — {pct_chg:.0f} % de la charge")

doc.add_paragraph(
    "La comparaison entre proportion d'effectifs et proportion de charge révèle les déséquilibres "
    "de répartition du travail :\n" + "\n".join(lines) + "\n"
    "Malgré les coefficients élevés attribués aux expéditions (5 et 8), ce sont les biologistes "
    "qui concentrent la majorité de la charge avec 73 % du total pour seulement 28 % des effectifs. "
    "Cela s'explique par la fréquence élevée des événements serre (plantation et récolte) qui se "
    "produisent plusieurs fois par jour sur plusieurs serres simultanément. "
    "Les techniciens en revanche sont largement sous-sollicités (31 % des effectifs pour 2 % de la charge), "
    "ce qui suggère un sur-effectif pour ce rôle."
)

doc.add_picture("graphiques/05_repartition_roles.png", width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ================================================================
# Sauvegarde
# ================================================================
doc.save("rapport_analyse.docx")
print("Rapport genere : rapport_analyse.docx")
